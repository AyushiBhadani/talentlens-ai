"""TalentLens AI — Resume Parser Service
Supports PDF (PyMuPDF) and DOCX (python-docx) formats.
"""
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Parse resume file and return extracted text.
    Supports PDF and DOCX formats.
    """
    ext = Path(filename).suffix.lower()
    
    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF and DOCX are supported.")


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                text_parts.append(text)
        
        doc.close()
        full_text = "\n\n".join(text_parts)
        
        if not full_text.strip():
            raise ValueError("PDF appears to be empty or contains only images (scanned PDF).")
        
        return full_text.strip()
        
    except ImportError:
        raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")
    except Exception as e:
        logger.error(f"PDF parsing error: {e}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)
        
        if not full_text.strip():
            raise ValueError("DOCX appears to be empty.")
        
        return full_text.strip()
        
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX parsing error: {e}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def validate_file(file_bytes: bytes, filename: str, max_size_mb: int = 10) -> None:
    """Validate file size and type."""
    # Check size
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > max_size_mb:
        raise ValueError(f"File too large: {size_mb:.1f}MB. Maximum allowed: {max_size_mb}MB.")
    
    # Check extension
    ext = Path(filename).suffix.lower()
    if ext not in (".pdf", ".docx", ".doc"):
        raise ValueError(f"Invalid file type: {ext}. Only PDF and DOCX are supported.")
    
    # Basic magic byte check for PDF
    if ext == ".pdf" and not file_bytes.startswith(b"%PDF"):
        raise ValueError("File does not appear to be a valid PDF.")
